# import os
# import requests
# import fitz  # PyMuPDF for PDF processing
# from tqdm import tqdm
# from langchain_community.embeddings import OllamaEmbeddings
# from langchain_chroma import Chroma
# from langchain_groq import ChatGroq
# from langchain.chains import RetrievalQA
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from crewai_tools import tool
# from dotenv import load_dotenv
# import pickle
# from docx import Document
# import io
# from flask import Flask, request, jsonify

# load_dotenv()

# app = Flask(__name__)


# @app.route('/setup_repo', methods=['POST'])
# def setup_repo():
#     data = request.json
#     repo_name = data.get('repo_name')
#     repo_url = data.get('repo_url')

#     # Call the setup function with the repository URL
#     message = setup_github_rag_system(repo_url)

#     return jsonify({'message': message})


# # Base path for persistence and processed files
# BASE_PATH = r"C:\Users\HP\Desktop\crew ai\project2\ragbase\db"


# def create_repo_folder(repo_url: str) -> str:
#     """
#     Create a directory for storing data related to the specific repository.

#     Parameters:
#     - repo_url: The GitHub repository URL.

#     Returns:
#     - The path of the created folder.
#     """
#     # Extract repository name from URL
#     repo_name = repo_url.rstrip("/").split("/")[-1]
#     # Create folder path based on repository name
#     folder_path = os.path.join(BASE_PATH, repo_name)
#     os.makedirs(folder_path, exist_ok=True)
#     return folder_path


# # @tool
# def setup_github_rag_system(
#     repo_url: str,
#     groq_api_key=os.getenv("GROQ_API_KEY"),
#     groq_model="llama3-8b-8192",
# ):
#     """
#     Set up the RAG system by loading files from a user-provided GitHub repo,
#     creating embeddings, storing them in a vector store, and setting up the QA chain.

#     Parameters:
#     - repo_url: The GitHub repository URL.
#     - groq_api_key: The API key for Groq.
#     - groq_model: The model to use with Groq (default is "llama3-8b-8192").

#     Returns:
#     - A string message indicating the setup completion.
#     """

#     if not groq_api_key:
#         return "Groq API key must be provided."

#     # Create a separate folder for the repository and adjust persistence directory
#     PERSIST_DIRECTORY = create_repo_folder(repo_url)
#     PROCESSED_FILES_PATH = os.path.join(
#         PERSIST_DIRECTORY, "processed_files.pkl")

#     # Load and process files from the GitHub repository
#     contents = get_github_repo_contents(repo_url)

#     if not os.path.exists(PERSIST_DIRECTORY):
#         print("Creating vector store and processing all files...")
#         vectorstore = create_new_vectorstore(
#             contents, PERSIST_DIRECTORY, PROCESSED_FILES_PATH
#         )
#     else:
#         print("Loading existing vector store...")
#         vectorstore = Chroma(
#             persist_directory=PERSIST_DIRECTORY,
#             embedding_function=OllamaEmbeddings(model="all-minilm:latest"),
#         )
#         print("Updating vector store with new files...")
#         update_vectorstore(vectorstore, contents, PROCESSED_FILES_PATH)

#     # Set up the retriever and QA chain
#     retriever = vectorstore.as_retriever()
#     llm = ChatGroq(api_key=groq_api_key, model_name=groq_model)
#     qa_chain = RetrievalQA.from_chain_type(
#         llm=llm, chain_type="stuff", retriever=retriever
#     )

#     # Store the objects globally for reuse
#     global stored_vectorstore, stored_qa_chain
#     stored_vectorstore = vectorstore
#     stored_qa_chain = qa_chain

#     return f"RAG system setup complete for repository: {repo_url}"


# def get_github_repo_contents(repo_url: str):
#     """Get the contents of a GitHub repository."""
#     repo_url = repo_url.rstrip("/")
#     api_url = (
#         repo_url.replace("https://github.com/",
#                          "https://api.github.com/repos/")
#         + "/contents/"
#     )
#     response = requests.get(api_url)
#     response.raise_for_status()
#     return response.json()


# def download_github_file(file_url: str) -> bytes:
#     """Download a file from a GitHub repository."""
#     response = requests.get(file_url)
#     response.raise_for_status()
#     return response.content


# def process_docx(file_content: bytes) -> str:
#     """Extract text from a .docx file."""
#     doc = Document(io.BytesIO(file_content))
#     text = "\n".join([p.text for p in doc.paragraphs])
#     return text


# def extract_text_from_pdf(file_content: bytes) -> str:
#     """Extract text from a PDF file using PyMuPDF (fitz)."""
#     pdf_document = fitz.open(stream=file_content, filetype="pdf")
#     text = ""
#     for page_num in range(pdf_document.page_count):
#         page = pdf_document.load_page(page_num)
#         text += page.get_text()
#     return text


# def process_github_files(contents, processed_files):
#     """Process and split files from the GitHub repository."""
#     documents = []

#     for item in tqdm(contents, desc="Processing files"):
#         if item["type"] == "file" and item["name"] not in processed_files:
#             file_url = item["download_url"]
#             file_name = item["name"]

#             try:
#                 if file_name.endswith((".md", ".txt")):
#                     file_content = download_github_file(
#                         file_url).decode("utf-8")
#                     documents.append(
#                         {"page_content": file_content, "type": "text"})

#                 elif file_name.endswith(".pdf"):
#                     file_content = download_github_file(file_url)
#                     text = extract_text_from_pdf(file_content)
#                     documents.append({"page_content": text, "type": "text"})

#                 elif file_name.endswith(".docx"):
#                     file_content = download_github_file(file_url)
#                     text = process_docx(file_content)
#                     documents.append({"page_content": text, "type": "text"})

#                 elif file_name.endswith(
#                     (
#                         ".py",
#                         ".js",
#                         ".java",
#                         ".cpp",
#                         ".c",
#                         ".html",
#                         ".css",
#                         ".sh",
#                         ".json",
#                         ".ts",
#                     )
#                 ):
#                     file_content = download_github_file(
#                         file_url).decode("utf-8")
#                     documents.append(
#                         {"page_content": file_content, "type": "code"})

#                 else:
#                     print(f"Skipping unsupported file type: {file_name}")
#             except UnicodeDecodeError:
#                 print(f"Skipping file {file_name} due to encoding issues.")
#             except Exception as e:
#                 print(f"An error occurred while processing {file_name}: {e}")

#     print(f"Processed {len(documents)} documents.")
#     return documents


# def create_new_vectorstore(contents, PERSIST_DIRECTORY, PROCESSED_FILES_PATH):
#     """Create a new vector store from the GitHub contents."""
#     processed_files = set()
#     documents = process_github_files(contents, processed_files)

#     if not documents:
#         print("No documents to process.")
#         return Chroma(
#             persist_directory=PERSIST_DIRECTORY,
#             embedding_function=OllamaEmbeddings(model="all-minilm:latest"),
#         )

#     print("Splitting text...")
#     text_splitter = RecursiveCharacterTextSplitter(
#         chunk_size=1000, chunk_overlap=200)

#     splits = []
#     for doc in documents:
#         split_texts = text_splitter.split_text(doc["page_content"])
#         for text in split_texts:
#             splits.append({"page_content": text})

#     print("Creating embeddings and vector store...")
#     vectorstore = create_embeddings_and_vectorstore(splits, PERSIST_DIRECTORY)

#     with open(PROCESSED_FILES_PATH, "wb") as f:
#         pickle.dump(
#             set([item["name"] for item in contents if item["type"] == "file"]), f
#         )

#     return vectorstore


# def update_vectorstore(vectorstore, contents, PROCESSED_FILES_PATH):
#     """Update an existing vector store with new GitHub contents."""
#     processed_files = set()
#     if os.path.exists(PROCESSED_FILES_PATH):
#         with open(PROCESSED_FILES_PATH, "rb") as f:
#             processed_files = pickle.load(f)

#     new_documents = process_github_files(contents, processed_files)

#     if new_documents:
#         print(f"Found {len(new_documents)} new documents to process.")

#         print("Splitting text for new documents...")
#         text_splitter = RecursiveCharacterTextSplitter(
#             chunk_size=1000, chunk_overlap=200
#         )
#         splits = []
#         for doc in new_documents:
#             split_texts = text_splitter.split_text(doc["page_content"])
#             for text in split_texts:
#                 splits.append({"page_content": text})

#         print("Creating embeddings for new documents...")
#         for split in tqdm(splits, desc="Adding new embeddings"):
#             vectorstore.add_texts(texts=[split["page_content"]])

#         # Update the processed files set and save it
#         new_processed_files = set(
#             item["name"] for item in contents if item["type"] == "file"
#         )
#         processed_files.update(new_processed_files)
#         with open(PROCESSED_FILES_PATH, "wb") as f:
#             pickle.dump(processed_files, f)
#     else:
#         print("No new documents found.")


# def create_embeddings_and_vectorstore(splits, persist_directory: str):
#     """Create embeddings and store them in a vector store."""
#     embedding_function = OllamaEmbeddings(model="all-minilm:latest")

#     # Ensure the directory exists
#     os.makedirs(persist_directory, exist_ok=True)

#     vectorstore = Chroma(
#         persist_directory=persist_directory,
#         embedding_function=embedding_function,
#     )

#     # We should pass the actual text content to add_texts
#     for split in tqdm(splits, desc="Creating embeddings and vector store"):
#         page_content = split.get("page_content")
#         if page_content:  # Ensure the content is not None or empty
#             vectorstore.add_texts(texts=[page_content])
#         else:
#             print("Skipping document with empty content")

#     return vectorstore


# # @tool
# def query_github_rag_system(query_text: str, file_name: str = None) -> str:
#     """
#     Query the documents from the GitHub RAG system.

#     Parameters:
#     - query_text: The question to ask about the content.
#     - file_name: (Optional) Specific file to query.

#     Returns:
#     - The answer to the query based on the content.
#     """

#     if "stored_qa_chain" not in globals() or "stored_vectorstore" not in globals():
#         return "RAG system not initialized. Please run setup_github_rag_system() first."

#     # Modify the query to focus on a specific file if file_name is provided
#     if file_name:
#         query_text = f"{query_text} in file {file_name}"

#     # Use invoke with the correct parameter name
#     answer = stored_qa_chain.invoke(input=query_text)
#     return answer


# if __name__ == "__main__":
#     repo_url = input("Enter your GitHub repository URL: ").strip()
#     setup_github_rag_system(repo_url)
#     print("RAG system setup complete. You can now ask questions.")

#     while True:
#         query_text = input(
#             "Enter your query (or type 'exit' to quit): ").strip()
#         if query_text.lower() == "exit":
#             print("Exiting the program.")
#             break

#         result = query_github_rag_system(query_text)
#         print(f"Query Result: {result}")


import os
import requests
import fitz  # PyMuPDF for PDF processing
from tqdm import tqdm
from langchain_community.embeddings import OllamaEmbeddings
from langchain_chroma import Chroma
from langchain_groq import ChatGroq
from langchain.chains import RetrievalQA
from langchain.text_splitter import RecursiveCharacterTextSplitter
from crewai_tools import tool
from dotenv import load_dotenv
import pickle
from docx import Document
import io
from flask import Flask, request, jsonify

load_dotenv()

app = Flask(__name__)


@app.route('/setup_repo', methods=['POST'])
def setup_repo():
    data = request.json
    repo_url = data.get('repo_url')

    # Call the setup function with the repository URL
    message = setup_github_rag_system(repo_url)

    return jsonify({'message': message})


# Base path for persistence and processed files
BASE_PATH = r"C:\Users\HP\Desktop\crew ai\project2\ragbase\db"


def create_repo_folder(repo_url: str) -> str:
    """Create a directory for storing data related to the specific repository."""
    repo_name = repo_url.rstrip("/").split("/")[-1]
    folder_path = os.path.join(BASE_PATH, repo_name)
    os.makedirs(folder_path, exist_ok=True)
    return folder_path


def setup_github_rag_system(
    repo_url: str,
    groq_api_key=os.getenv("GROQ_API_KEY"),
    groq_model="llama3-8b-8192",
):
    """Set up the RAG system by processing all files from a user-provided GitHub repo."""

    if not groq_api_key:
        return "Groq API key must be provided."

    # Create a separate folder for the repository and adjust persistence directory
    PERSIST_DIRECTORY = create_repo_folder(repo_url)
    PROCESSED_FILES_PATH = os.path.join(
        PERSIST_DIRECTORY, "processed_files.pkl")

    # Load and process files from the GitHub repository
    contents = get_github_repo_contents(repo_url)

    if not os.path.exists(PERSIST_DIRECTORY):
        print("Creating vector store and processing all files...")
        vectorstore = create_new_vectorstore(
            contents, PERSIST_DIRECTORY, PROCESSED_FILES_PATH
        )
    else:
        print("Loading existing vector store...")
        vectorstore = Chroma(
            persist_directory=PERSIST_DIRECTORY,
            embedding_function=OllamaEmbeddings(model="all-minilm:latest"),
        )
        print("Updating vector store with new files...")
        update_vectorstore(vectorstore, contents, PROCESSED_FILES_PATH)

    # Set up the retriever and QA chain
    retriever = vectorstore.as_retriever()
    llm = ChatGroq(api_key=groq_api_key, model_name=groq_model)
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm, chain_type="stuff", retriever=retriever
    )

    global stored_vectorstore, stored_qa_chain
    stored_vectorstore = vectorstore
    stored_qa_chain = qa_chain

    return f"RAG system setup complete for repository: {repo_url}"


def get_github_repo_contents(repo_url: str):
    """Get the contents of a GitHub repository, including directories."""
    repo_url = repo_url.rstrip("/")
    api_url = (
        repo_url.replace("https://github.com/",
                         "https://api.github.com/repos/")
        + "/contents/"
    )
    contents = []
    get_contents_recursive(api_url, contents)
    return contents


def get_contents_recursive(api_url, contents):
    """Recursively retrieve all files from the GitHub repository, including directories."""
    response = requests.get(api_url)
    response.raise_for_status()
    items = response.json()
    for item in items:
        if item["type"] == "file":
            contents.append(item)
        elif item["type"] == "dir":
            get_contents_recursive(item["url"], contents)


def download_github_file(file_url: str) -> bytes:
    """Download a file from a GitHub repository."""
    response = requests.get(file_url)
    response.raise_for_status()
    return response.content


def process_github_files(contents, processed_files):
    """Process and split files from the GitHub repository."""
    documents = []

    for item in tqdm(contents, desc="Processing files"):
        if item["name"] not in processed_files:
            file_url = item["download_url"]

            try:
                file_content = download_github_file(file_url)

                # Determine file type and process accordingly
                if file_url.endswith(".pdf"):
                    # Handle PDF files
                    pdf_document = fitz.open(
                        stream=file_content, filetype="pdf")
                    text = ""
                    for page in pdf_document:
                        text += page.get_text()
                    documents.append({"page_content": text})
                elif file_url.endswith(".docx"):
                    # Handle DOCX files
                    doc = Document(io.BytesIO(file_content))
                    text = "\n".join([p.text for p in doc.paragraphs])
                    documents.append({"page_content": text})
                else:
                    # Assume text file
                    file_content = file_content.decode(
                        "utf-8", errors="ignore")
                    documents.append({"page_content": file_content})

            except Exception as e:
                print(f"An error occurred while processing {
                      item['name']}: {e}")

    print(f"Processed {len(documents)} documents.")
    return documents


def create_new_vectorstore(contents, PERSIST_DIRECTORY, PROCESSED_FILES_PATH):
    """Create a new vector store from the GitHub contents."""
    processed_files = set()
    documents = process_github_files(contents, processed_files)

    if not documents:
        print("No documents to process.")
        return Chroma(
            persist_directory=PERSIST_DIRECTORY,
            embedding_function=OllamaEmbeddings(model="all-minilm:latest"),
        )

    print("Splitting text...")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000, chunk_overlap=200)

    splits = []
    for doc in documents:
        split_texts = text_splitter.split_text(doc["page_content"])
        for text in split_texts:
            splits.append({"page_content": text})

    print("Creating embeddings and vector store...")
    vectorstore = create_embeddings_and_vectorstore(splits, PERSIST_DIRECTORY)

    with open(PROCESSED_FILES_PATH, "wb") as f:
        pickle.dump(set([item["name"] for item in contents]), f)

    return vectorstore


def update_vectorstore(vectorstore, contents, PROCESSED_FILES_PATH):
    """Update an existing vector store with new GitHub contents."""
    processed_files = set()
    if os.path.exists(PROCESSED_FILES_PATH):
        with open(PROCESSED_FILES_PATH, "rb") as f:
            processed_files = pickle.load(f)

    new_documents = process_github_files(contents, processed_files)

    if new_documents:
        print(f"Found {len(new_documents)} new documents to process.")

        print("Splitting text for new documents...")
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000, chunk_overlap=200)
        splits = []
        for doc in new_documents:
            split_texts = text_splitter.split_text(doc["page_content"])
            for text in split_texts:
                splits.append({"page_content": text})

        print("Creating embeddings for new documents...")
        for split in tqdm(splits, desc="Adding new embeddings"):
            vectorstore.add_texts(texts=[split["page_content"]])

        new_processed_files = set(item["name"] for item in contents)
        processed_files.update(new_processed_files)
        with open(PROCESSED_FILES_PATH, "wb") as f:
            pickle.dump(processed_files, f)
    else:
        print("No new documents found.")


def create_embeddings_and_vectorstore(splits, persist_directory: str):
    """Create embeddings and store them in a vector store."""
    embedding_function = OllamaEmbeddings(model="all-minilm:latest")

    os.makedirs(persist_directory, exist_ok=True)

    vectorstore = Chroma(
        persist_directory=persist_directory,
        embedding_function=embedding_function,
    )

    for split in tqdm(splits, desc="Creating embeddings and vector store"):
        page_content = split.get("page_content")
        if page_content:
            vectorstore.add_texts(texts=[page_content])
        else:
            print("Skipping document with empty content")

    return vectorstore


def query_github_rag_system(query_text: str) -> str:
    """Query the documents from the GitHub RAG system."""

    if "stored_qa_chain" not in globals() or "stored_vectorstore" not in globals():
        return "RAG system not initialized. Please run setup_github_rag_system() first."

    answer = stored_qa_chain.invoke(input=query_text)
    return answer


if __name__ == "__main__":
    repo_url = input("Enter your GitHub repository URL: ").strip()
    setup_github_rag_system(repo_url)
    print("RAG system setup complete. You can now ask questions.")

    while True:
        query_text = input(
            "Enter your query (or type 'exit' to quit): ").strip()
        if query_text.lower() == "exit":
            print("Exiting the program.")
            break

        result = query_github_rag_system(query_text)
        print(f"Query Result: {result}")
